from django.shortcuts import render
from django.http import HttpResponse
from .algorithm import main
from docx import Document
from .algorithm import fileSimilarity
import PyPDF2
import os
import tempfile

# Create your views here.

# Home page
def home(request):
    return render(request, 'pc/index.html')

# Web search (Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])
    
    if request.POST['q']:
        percent, link = main.findSimilarity(request.POST['q'])
        percent = round(percent, 2)
    print("Output.....................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Web search file (.txt, .docx, .pdf) - FIXED PDF HANDLING
def filetest(request):
    value = ''
    print(request.FILES['docfile'])
    
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read(), 'utf-8')
    
    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text
    
    elif str(request.FILES['docfile']).endswith(".pdf"):
        # FIXED: Correctly read PDF file from uploaded file object
        pdf_reader = PyPDF2.PdfReader(request.FILES['docfile'])
        value = ""
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                value += extracted_text
    
    percent, link = main.findSimilarity(value)
    print("Output...................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Text compare page
def fileCompare(request):
    return render(request, 'pc/doc_compare.html')

# Two text compare (Text)
def twofiletest1(request):
    print("Submitted text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])
    
    result = 0
    if request.POST['q1'] != '' and request.POST['q2'] != '':
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'], request.POST['q2'])
        result = round(result, 2)
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})

# Two file compare (.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    
    if str(request.FILES['docfile1']).endswith(".txt") and str(request.FILES['docfile2']).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read(), 'utf-8')
        value2 = str(request.FILES['docfile2'].read(), 'utf-8')
    
    elif str(request.FILES['docfile1']).endswith(".docx") and str(request.FILES['docfile2']).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text
    
    result = fileSimilarity.findFileSimilarity(value1, value2)
    result = round(result, 2)
    
    print("Output..................!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})
